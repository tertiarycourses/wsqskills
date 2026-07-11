#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for a course: a Written Assessment (SAQ) that tests KNOWLEDGE,
plus ONE practical instrument that tests PRACTICAL ability — either a Case Study (CS) or a
Practical Performance (PP), selected with the INSTRUMENT constant below.

FOLLOW THE ORIGINAL: if the course already has a reference/previous assessment, keep its
instrument type (CS or PP), its question/task COUNT, its criterion codes (K.. / A..) and its
timings. Never migrate a Case Study course to PP, or the reverse.

Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Certified Lean Six Sigma Yellow Belt (CLSSYB) Training"
COURSE_CODE = "TGS-2025053922"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = None   # Tertiary-only cover (as per the CLSSYB Lesson Plan / Learner Guide)

# INSTRUMENT — the practical instrument this course is assessed with. It MUST match the course's
# ORIGINAL / reference assessment: "CS" = Case Study, "PP" = Practical Performance.
# NEVER convert a Case Study course to PP (or the reverse) — see SKILL.md, "Follow the original".
INSTRUMENT = "CS"           # CLSSYB reference assessment is WA (SAQ) + Case Study
WA_MINUTES, PRACTICAL_MINUTES = "60 minutes", "90 minutes"   # taken from the reference papers
# ANSWER_STYLE — how the non-table lines of a model answer are rendered.
#   "code"  → Consolas, indentation preserved (commands / Dockerfiles / YAML: Docker, K8s, n8n …)
#   "prose" → Arial that wraps naturally (written/management courses: Lean Six Sigma, PM …)
# Pipe-delimited rows always become real Word tables regardless of this setting.
ANSWER_STYLE = "prose"

Q_VER, A_VER = "v1", "v1"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ------------------------------------------------- COMPETENCY UNIT (coverage contract)
# Codes follow the reference CLSSYB assessment: WA = K1–K2, PP = A1–A5.
# A1  Explain Yellow Belt support responsibilities
# A2  Identify customer value, waste, defects and variation
# A3  Use SIPOC, PDCA and DMAIC Define tools
# A4  Plan data collection and perform basic analysis
# A5  Support root cause, countermeasures and control planning
KNOWLEDGE = ["K1", "K2"]                          # ← tested by WRITTEN (3 questions)
ABILITIES = ["A1", "A2", "A3", "A4", "A5"]        # ← tested by PRACTICAL (3 tasks)

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 ("K1",
  "Lean improves a process by removing everything the customer would not be willing to pay for. In a service "
  "process such as an IT service desk or a clinic, that wasted effort is usually hidden in queues, rework and "
  "chasing information rather than in physical stock.",
  "Explain the meaning of the term “waste” in Lean, and list and briefly describe three common types of waste "
  "that can occur in a healthcare or service environment.",
  ["Waste (Japanese: muda) is any activity that consumes time, effort, people or money but adds no value in the "
   "eyes of the customer — the customer would not be willing to pay for it. Lean removes waste so that the "
   "value-adding steps flow.",
   "Lean classifies work as value-added (changes the service/product in a way the customer values), "
   "business-value-added (needed for compliance or control) or non-value-added (pure waste, to be eliminated).",
   "Award the mark for any three of the eight wastes, correctly described with a service example. For instance:",
   "Waiting — the customer or the work sits idle: a ticket waits in the queue before it is assigned; a patient "
   "waits for the doctor during the morning peak.",
   "Defects — work that is wrong and must be redone: a ticket logged in the wrong queue or with missing "
   "information, so it is reopened or re-routed.",
   "Extra processing — doing more than the customer needs: re-keying the same details into two systems, or "
   "sending repeated status updates that carry no new information.",
   "Other acceptable answers: overproduction (reports nobody reads), transportation (passing a case between "
   "departments), motion (staff walking or clicking between screens), inventory (a growing ticket backlog), "
   "non-utilised talent (skilled staff doing clerical work).",
   "(Slides: Lean, Six Sigma and Value / The Eight Wastes — Lab 2, waste walk.)"]),
 ("K1, K2",
  "DMAIC is the improvement roadmap used on larger Six Sigma projects. A Yellow Belt does not usually lead the "
  "project, but is expected to support the early phases with process knowledge and data.",
  "In the context of the DMAIC methodology used in Lean Six Sigma, describe what activities are typically "
  "carried out in the Define and Measure phases, and provide one tool used in each phase and explain its purpose.",
  ["Define phase — agree what problem is being solved before any solution is discussed: write the problem "
   "statement (process, time period, measurable issue, impact — with no assumed cause and no assumed solution) "
   "and the goal statement; set the scope (start point, end point, what is in and out); capture the Voice of the "
   "Customer and translate it into CTQ requirements; identify stakeholders and the team; state the expected "
   "business benefit.",
   "Define tool — SIPOC (Suppliers, Inputs, Process, Outputs, Customers): a one-page, high-level view of the "
   "process (five to seven steps) whose purpose is to fix the process boundaries and identify who supplies and "
   "who receives the work, so the team agrees on scope before detailed mapping. (A Project Charter is equally "
   "acceptable — purpose: it records problem, goal, scope, benefit, timeline and team in one authorised document.)",
   "Measure phase — establish the baseline with facts, not opinion: choose KPIs and write an operational "
   "definition for each; plan the sampling (sample size, period, who records, how data is verified, likely bias); "
   "collect the data; calculate the baseline metrics (defect rate, average and range of cycle time, yield, DPU, "
   "DPMO).",
   "Measure tool — a check sheet: a simple structured form (date, ticket ID, type, assignment time, defect "
   "category, rework, notes) whose purpose is to make data capture consistent and easy at the point of work so "
   "every recorder counts the same thing the same way. (A data collection plan or run chart is equally acceptable.)",
   "The candidate should show that Define answers “what is the problem?” and Measure answers “how big is it "
   "today?” — neither phase proposes a solution.",
   "(Slides: DMAIC Overview / Define / Measure — Labs 3, 5 and 6.)"]),
 ("K2",
  "A service desk (or a clinic) is receiving frequent complaints about delays. The team is under pressure and "
  "the first instinct is to add more staff.",
  "Describe how you would use the Fishbone (Ishikawa) Diagram and the 5 Whys technique to identify possible root "
  "causes of this problem, and explain why it is important to identify the root cause rather than just "
  "addressing the symptoms.",
  ["State the problem precisely first, as a single focused statement — e.g. “Ticket assignment is delayed for "
   "30% of tickets during the morning peak, causing repeat status calls and missed service levels.”",
   "Fishbone (Ishikawa) — write the problem in the fish head, then brainstorm possible causes under each "
   "category branch: People, Process, Policy, Technology, Measurement and Environment (the Man / Method / "
   "Machine / Material / Measurement / Environment 6M form is equally acceptable). It is used to widen thinking "
   "so the team sees every possible cause instead of the first one that comes to mind.",
   "Typical branches for service delay: People — not enough triage staff at peak, new agents untrained. Process "
   "— no triage checklist, unclear ownership after handoff. Policy — no rule on how fast a ticket must be "
   "assigned. Technology — no auto-routing, free-text categories. Measurement — “resolved” defined "
   "differently by each agent. Environment — arrivals all bunch at the start of the day.",
   "5 Whys — take one likely cause from the fishbone and ask “why?” repeatedly (about five times), each answer "
   "becoming the next question, until you reach a cause that can actually be tested or controlled. Example: "
   "tickets are assigned late → why? they sit unassigned in the general queue → why? nobody owns triage at peak "
   "→ why? triage is not in anyone's standard work → why? the category field is free text so tickets cannot be "
   "auto-routed → root cause: no standard categories / no auto-routing rule.",
   "Record the evidence needed at each level (data, observation, interview) and rate each candidate cause "
   "high/medium/low confidence — a cause is only a root cause once evidence supports it.",
   "Why root cause and not symptoms: treating the symptom (adding staff, chasing tickets, apologising) gives "
   "only temporary relief — the problem returns as soon as the workaround stops, and it wastes effort and cost. "
   "Fixing the root cause removes the source of the defect permanently, so the improvement can be standardised "
   "and held with a control plan.",
   "(Slides: Analyze — Root Cause Analysis / 5 Whys / Fishbone — Lab 8.)"]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "Contoso Service Desk handles employee IT requests. Employees complain that tickets take too long to be "
 "assigned, that they must keep chasing for status, and that agents apply different definitions of “resolved”. "
 "Handoffs between the front line and the specialist queues are unclear, and although ticket data exists it is "
 "inconsistent. Two weeks of data have been collected on 400 tickets: 96 tickets contained at least one defect, "
 "and 120 defects were recorded in total — Delayed assignment 48, Missing information 30, Wrong queue 18, "
 "Reopened ticket 12, Duplicate ticket 8, Unclear status 4. Each ticket has 6 defect opportunities. "
 "You are the Lean Six Sigma Yellow Belt supporting this team. You are not leading the project: you support the "
 "team by mapping the process, collecting and presenting data, helping to find the root cause, and running a "
 "small PDCA improvement under guidance. Complete the three tasks below — each mirrors a hands-on lab you did "
 "in class. You may use your completed lab worksheets as evidence.")

# (label, criterion, task prompt, box caption, model-answer build steps citing the activity)
BOX_CAP = "Write your answer (tables, diagram and calculations) in the box below"
PRACTICAL = [
 ("Task 1", "A1, A2",
  "Yellow Belt role, Voice of the Customer, and waste (Define the value). "
  "Part A — Produce a table of at least four Yellow Belt responsibilities on this project (Responsibility | Your "
  "contribution), and state in one or two sentences how the Yellow Belt role differs from that of a Green Belt "
  "or Black Belt. "
  "Part B — Build a Voice of the Customer table with at least three rows (Customer statement | Need | CTQ "
  "requirement) using the employee complaints above. "
  "Part C — Carry out a waste walk on the ticket process: identify at least three of the eight wastes, with the "
  "evidence for each, and classify at least three process activities as value-added, business-value-added or "
  "non-value-added. "
  "Part D — For this process, state what counts as a defect, what is an opportunity for a defect, and what "
  "variation is visible in the data. (Labs 1 and 2.)",
  BOX_CAP,
  "Part A — Yellow Belt role (Lab 1, Step 1):\n"
  "  Responsibility            | Yellow Belt contribution\n"
  "  Basic LSS knowledge       | Applies Lean/Six Sigma vocabulary and tools correctly in the team\n"
  "  Process map support       | Builds the SIPOC and the detailed map of ticket assignment\n"
  "  Subject matter expert     | Explains what really happens, where work sticks, which rules are informal\n"
  "  Data collection support   | Designs the check sheet, collects and verifies the ticket data\n"
  "  PDCA small improvement    | Runs the triage-checklist pilot under the Green Belt's guidance\n"
  "  Team participation        | Attends reviews, presents findings, supports the project leader\n"
  "  Role difference: a Yellow Belt SUPPORTS projects and may run small PDCA improvements; a Green Belt leads a\n"
  "  full DMAIC project part-time; a Black Belt leads large, cross-functional projects full-time and coaches\n"
  "  others. CSSC sets no prerequisite and no project requirement for Yellow Belt.\n"
  "\n"
  "Part B — Voice of the Customer (Lab 2, Step 1):\n"
  "  Customer statement                          | Need                | CTQ requirement\n"
  "  \"I don't know when my ticket will be handled\" | Status visibility  | Status update within 1 h of logging\n"
  "  \"It takes too long before anyone picks it up\"  | Fast assignment    | Ticket assigned within 30 min\n"
  "  \"They closed it but it wasn't fixed\"           | Consistent closure | Single agreed definition of 'resolved'\n"
  "  Accept any reasonable statement/need/CTQ set, provided the CTQ is specific and measurable.\n"
  "\n"
  "Part C — Waste walk and value analysis (Lab 2, Steps 2–3):\n"
  "  Waiting          | Tickets sit unassigned in the general queue (the largest defect category, 48)\n"
  "  Defects          | Wrong queue (18) and missing information (30) force re-routing and rework\n"
  "  Extra processing | Agents re-key details; employees send repeat status chasers\n"
  "  Inventory        | Backlog of unassigned tickets builds up at peak\n"
  "  Non-utilised talent | Specialists spend time re-categorising tickets instead of resolving them\n"
  "  Value analysis: diagnosing/resolving the issue = value-added; logging and audit trail =\n"
  "  business-value-added; re-routing, chasing status and re-keying = non-value-added.\n"
  "\n"
  "Part D — Defects, opportunities and variation (Lab 2, Step 4):\n"
  "  Defect: any ticket outcome that fails the CTQ — delayed assignment, missing information, wrong queue,\n"
  "    duplicate, reopened, or unclear status.\n"
  "  Opportunity for a defect: each of the 6 points in handling a ticket where such a defect can occur\n"
  "    (6 opportunities per ticket).\n"
  "  Variation: assignment time differs from ticket to ticket and from day to day (peaks at start of day);\n"
  "    'resolved' is applied inconsistently between agents. Data would prove it: assignment-time run chart.\n"
  "  Award the mark where the candidate distinguishes waste (no value) from a defect (work done wrong)."),
 ("Task 2", "A3, A4",
  "Map the process, charter the improvement, and analyse the data (Define → Measure → Analyze). "
  "Part A — Draw a SIPOC for the ticket assignment process (Suppliers | Inputs | Process | Outputs | Customers), "
  "keeping the process to five to seven high-level steps, and mark at least three pain points (waiting, rework "
  "loop, unclear ownership, duplicate entry or missing decision rule) on the flow. "
  "Part B — Write a PDCA improvement charter for a two-week pilot of a ticket triage checklist. It must contain: "
  "a problem statement (process, time period, measurable issue, impact — no assumed cause, no assumed solution), "
  "a goal statement (metric, baseline, target, date), scope (in / out), at least two stakeholders with their "
  "interest, and the Plan-Do-Check-Act actions with the decision rule to adopt, adapt or abandon. "
  "Part C — Produce a data collection plan: at least three KPIs with an operational definition, data source and "
  "frequency, plus the fields of the check sheet you would use. "
  "Part D — Using the two weeks of data in the scenario, build the Pareto table (category, count, percent, "
  "cumulative percent) sorted highest to lowest, calculate the defect rate, yield, DPU and DPMO, and state in "
  "two or three sentences which category the team should attack first and what conclusion cannot yet be drawn. "
  "(Labs 3, 4, 5, 6 and 7.)",
  BOX_CAP,
  "Part A — SIPOC (Lab 3, Steps 1–3):\n"
  "  Suppliers | Inputs            | Process (5–7 steps)      | Outputs          | Customers\n"
  "  Employee  | Request details   | 1 Log ticket             | Assigned ticket  | Employee\n"
  "  Helpdesk  | Category/priority | 2 Categorise             | Status update    | IT specialist team\n"
  "  IT system | Ticket template   | 3 Triage / prioritise    | Resolution note  | Service desk manager\n"
  "  Manager   | Routing rules     | 4 Assign to queue        | Closure record   |\n"
  "            | Agent availability| 5 Resolve                | KPI data         |\n"
  "            |                   | 6 Confirm and close      |                  |\n"
  "  Pain points to mark: WAITING before assignment (step 3→4); REWORK loop when the queue is wrong (5→2);\n"
  "  UNCLEAR OWNERSHIP after handoff (nobody owns triage at peak); DUPLICATE ENTRY (details re-keyed);\n"
  "  MISSING DECISION RULE (no stated time limit for assignment).\n"
  "\n"
  "Part B — PDCA charter (Labs 4 and 5):\n"
  "  Problem: Over the last two weeks, the ticket assignment process delayed 48 of 400 tickets beyond the\n"
  "    30-minute target, causing repeat status chasers and missed service levels.\n"
  "  Goal: Reduce average ticket assignment time from 55 minutes (baseline) to under 30 minutes by <date>,\n"
  "    and cut delayed assignments by half.\n"
  "  Scope — In: logging to assignment of employee IT tickets at the Contoso Service Desk.\n"
  "         Out: resolution work inside specialist queues; procurement and vendor tickets.\n"
  "  Stakeholders: service desk manager (owns the KPI, needs weekly figures); triage agents (do the work,\n"
  "    need the checklist to be quick); employees/customers (want status visibility); Green Belt project lead.\n"
  "  PDCA: Plan — design the triage checklist and the measurement; Do — pilot it with one triage team for\n"
  "    two weeks; Check — compare assignment time and defect counts against the baseline; Act — decision rule:\n"
  "    ADOPT and standardise if assignment time falls below 30 min; ADAPT if it improves but misses target;\n"
  "    ABANDON if there is no improvement.\n"
  "  The problem statement must NOT name a cause or a solution — mark down if it does.\n"
  "\n"
  "Part C — Data collection plan (Lab 6, Steps 1–4):\n"
  "  KPI                 | Operational definition                          | Source        | Frequency\n"
  "  Assignment time     | Minutes from ticket logged to ticket assigned   | Ticket system | Every ticket\n"
  "  First response time | Minutes from logged to first reply to employee  | Ticket system | Every ticket\n"
  "  Reopen rate         | Reopened tickets / closed tickets in the week   | Ticket system | Weekly\n"
  "  (Backlog of unassigned tickets and status-chaser count are equally acceptable.)\n"
  "  Check sheet fields: Date, Ticket ID, Ticket type, Assignment time, Defect category, Rework required, Notes.\n"
  "  Sampling: all tickets over two weeks; recorded by the triage agent; verified weekly against the\n"
  "  system export; bias risk — agents may under-record their own delays.\n"
  "\n"
  "Part D — Pareto and metrics (Lab 7, Steps 1–5):\n"
  "  Defect category | Count | Percent | Cumulative %\n"
  "  Delayed assignment | 48 | 40.0% | 40.0%\n"
  "  Missing information | 30 | 25.0% | 65.0%\n"
  "  Wrong queue | 18 | 15.0% | 80.0%\n"
  "  Reopened ticket | 12 | 10.0% | 90.0%\n"
  "  Duplicate ticket | 8 | 6.7% | 96.7%\n"
  "  Unclear status | 4 | 3.3% | 100.0%\n"
  "  Total | 120 |  | \n"
  "\n"
  "  Defect rate  = 96 defective tickets / 400 = 24%\n"
  "  Yield        = (400 - 96) / 400 = 304/400 = 76%\n"
  "  DPU          = 120 defects / 400 units = 0.30\n"
  "  DPMO         = 120 / (400 x 6) x 1,000,000 = 50,000\n"
  "  Interpretation: attack Delayed assignment first — with Missing information it accounts for 65% of all\n"
  "  defects (the vital few). What cannot yet be concluded: WHY assignment is delayed — the Pareto shows where\n"
  "  the defects are, not their cause; a run chart of daily assignment time is needed to separate normal\n"
  "  variation from a real trend, and no single day's data point should be over-reacted to.\n"
  "  Accept minor rounding differences; the ranking, the 65% vital few, DPU 0.30 and DPMO 50,000 must be right."),
 ("Task 3", "A5",
  "Find the root cause, choose countermeasures, and hold the gain (Analyze → Improve → Control). "
  "Part A — Construct a Fishbone (Ishikawa) diagram for “Ticket assignment is delayed” using at least four "
  "categories (e.g. People, Process, Policy, Technology, Measurement, Environment), with at least two possible "
  "causes on each branch. "
  "Part B — Select one likely cause from your diagram and perform a 5 Whys analysis (Why level | Answer | "
  "Evidence needed) down to a cause that can be tested or controlled, then state what evidence would validate it. "
  "Part C — Propose at least three countermeasures for that root cause (Root cause | Countermeasure | Expected "
  "effect). At least one must be a mistake-proofing (poka-yoke) device and one must be standard work; apply 5S "
  "thinking to at least two of the five S's for this information process. "
  "Part D — Build a control plan so the gain is sustained (Process step | Metric | Target | Owner | Check "
  "frequency | Response plan) with at least three rows, name the visual management you would display, and list "
  "the handover items the process owner must accept. (Labs 8, 9 and 10.)",
  BOX_CAP,
  "Part A — Fishbone for “Ticket assignment is delayed” (Lab 8, Step 3):\n"
  "  People       | No owner for triage at peak; new agents untrained; leave/absence at 9am\n"
  "  Process      | No triage checklist; unclear handoff to specialist queues; no priority rule\n"
  "  Policy       | No stated time limit for assignment; 'resolved' not defined; no escalation trigger\n"
  "  Technology   | Category field is free text; no auto-routing rule; no queue alert\n"
  "  Measurement  | Assignment time not reported; agents count 'resolved' differently\n"
  "  Environment  | Tickets arrive in a burst at the start of the day; interruptions at the desk\n"
  "  (The 6M form — Man, Method, Machine, Material, Measurement, Environment — is equally acceptable.)\n"
  "\n"
  "Part B — 5 Whys on “Tickets sit unassigned in the general queue” (Lab 8, Steps 2 and 5):\n"
  "  Why 1 | Tickets are assigned late          | Evidence: assignment-time data, 48 delayed tickets\n"
  "  Why 2 | They sit in the general queue      | Evidence: queue-age report / observation at peak\n"
  "  Why 3 | Nobody owns triage during the peak | Evidence: roster, interview with agents\n"
  "  Why 4 | Triage is in nobody's standard work| Evidence: no documented triage procedure\n"
  "  Why 5 | Categories are free text, so tickets cannot be auto-routed and must be read one by one | Evidence: system config; 18 wrong-queue and 30 missing-information defects\n"
  "\n"
  "  ROOT CAUSE: no standard ticket categories and no auto-routing / triage ownership rule.\n"
  "  Stop at a cause the team can test or control. Validate with: system configuration check, two-week\n"
  "  queue-age data, observation of the morning peak, and agent interviews. Rate each candidate cause\n"
  "  high/medium/low confidence before acting.\n"
  "\n"
  "Part C — Countermeasures (Lab 9, Steps 1–5):\n"
  "  Root cause | Countermeasure | Expected effect\n"
  "  Free-text categories | POKA-YOKE: mandatory dropdown category and required fields before the ticket can be submitted | Wrong-queue and missing-information defects fall; auto-routing becomes possible\n"
  "  No triage ownership at peak | STANDARD WORK: triage checklist with a named owner per shift, 9:00–11:00 | Tickets assigned within 30 minutes\n"
  "  Delays are invisible | VISUAL MANAGEMENT: queue-age board with a 30-minute escalation trigger | Delays are seen and acted on before they breach\n"
  "\n"
  "  Prioritise with Impact – Effort – Risk scoring (1–5 each) and pick low-effort, low-risk, high-impact first.\n"
  "  5S applied to an information process: SORT — close duplicate and dead tickets; SET IN ORDER — one queue\n"
  "  per category with a clear naming rule; SHINE — weekly clean-up of stale tickets; STANDARDIZE — the triage\n"
  "  checklist and the single definition of 'resolved'; SUSTAIN — weekly KPI review with the process owner.\n"
  "\n"
  "Part D — Control plan (Lab 10, Steps 1–4):\n"
  "  Process step | Metric | Target | Owner | Check frequency | Response if control is lost\n"
  "  Triage | Assignment time | < 30 min | Triage lead | Daily | Escalate to the duty manager; re-staff the peak; find the cause\n"
  "  Categorise | Wrong-queue defects | < 2 / week | Service desk manager | Weekly | Review the dropdown and routing rule; retrain the agent\n"
  "  Log ticket | Checklist compliance | 100% | Triage lead | Weekly | Coach the agent; re-issue the standard work\n"
  "  Close | Reopen rate | < 5% | Service desk manager | Weekly | Re-check the single definition of 'resolved'\n"
  "\n"
  "  Visual management: weekly run chart of assignment time; defect-category Pareto; queue status board;\n"
  "  checklist-compliance count; escalation trigger list — displayed where the team can see them.\n"
  "  Handover: process owner accepts the control plan; standard work is stored and findable; a KPI report\n"
  "  owner is named; the review cadence is scheduled; open issues are listed; escalation criteria documented.\n"
  "  The A3 summary (background, current condition, goal, analysis, countermeasures, implementation, results,\n"
  "  follow-up/control) is the one-page handover document.\n"
  "  Award the mark where every control row has a metric, a target, an owner, a frequency AND a response plan —\n"
  "  a plan with no response action is not a control plan."),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def _mono_line(cell, text):
    """One non-table line of a model answer inside the answer box.

    ANSWER_STYLE = "code"  → fixed-width Consolas, indentation preserved. Correct for commands,
                             Dockerfiles and YAML, where alignment carries meaning.
    ANSWER_STYLE = "prose" → Arial that wraps naturally, with the source indent turned into a real
                             left indent. Correct for written courses: a long sentence then wraps
                             cleanly under its own indent instead of overflowing the fixed grid.
    """
    b = cell.add_paragraph(style=None)
    b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
    if ANSWER_STYLE == "prose":
        indent = len(text) - len(text.lstrip(" "))
        b.paragraph_format.left_indent = Inches(0.09 * (indent // 2))
        rr = b.add_run(text.strip() or " ")
        rr.font.name = "Arial"; rr.font.size = Pt(9.5)
        return
    rr = b.add_run(text if text else " ")
    rr.font.name = "Consolas"; rr.font.size = Pt(8.5)
    rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
    wt = rr._element.find(qn('w:t'))
    if wt is not None: wt.set(qn('xml:space'), 'preserve')

def _reflow(lines):
    """Undo the source's fixed-width hard wrap so prose becomes whole paragraphs.

    Model answers are authored ~100 chars wide for readability in the source. Rendered as Arial
    they wrap proportionally, so those hard breaks would leave mid-sentence orphan lines. A line
    is a CONTINUATION of the one above when it starts lower-case/punctuation, or is indented
    deeper — anything else (a new sentence, a 'Part B —' heading, a blank line) starts a paragraph.
    """
    out = []
    for ln in lines:
        s = ln.strip()
        if not s:
            out.append(ln); continue
        prev = out[-1] if out else ""
        cont = (prev.strip()
                and not prev.strip().endswith(":")
                and (not s[0].isupper()
                     or (len(ln) - len(ln.lstrip(" "))) > (len(prev) - len(prev.lstrip(" ")))))
        if cont:
            out[-1] = prev.rstrip() + " " + s
        else:
            out.append(ln)
    return out

def _grid_table(cell, rows):
    """Render pipe-delimited model-answer rows as a REAL Word table nested in the answer box.
    ASCII pipe-tables wrap and lose their column alignment once a row is wider than the text
    frame; a native table wraps INSIDE the cell, so the columns always line up."""
    grid = [[c.strip() for c in ln.split("|")] for ln in rows]
    ncol = max(len(r) for r in grid)
    grid = [r + [""] * (ncol - len(r)) for r in grid]
    # A first row whose cells are all short is a header row → bold it.
    header = all(len(c) <= 30 for c in grid[0])
    t = cell.add_table(rows=len(grid), cols=ncol); t.style = "Table Grid"; t.autofit = True
    for i, row in enumerate(grid):
        trPr = t.rows[i]._tr.get_or_add_trPr(); trPr.append(OxmlElement('w:cantSplit'))
        for j, val in enumerate(row):
            c = t.cell(i, j); c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
            rr = p.add_run(val); rr.font.name = "Arial"; rr.font.size = Pt(8.5)
            rr.bold = bool(header and i == 0)
    cell.add_paragraph().paragraph_format.space_after = Pt(2)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → model-answer block, where
    consecutive pipe-delimited rows become a real nested Word table and everything else stays
    monospace (commands, YAML, calculations); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        buf, prose = [], []
        def flush_prose():
            for ln in _reflow(prose) if ANSWER_STYLE == "prose" else prose:
                _mono_line(cell, ln)
            prose.clear()
        for ln in code.split("\n"):
            if "|" in ln:
                flush_prose()
                buf.append(ln); continue
            if buf:
                _grid_table(cell, buf); buf = []
            prose.append(ln)
        flush_prose()
        if buf:
            _grid_table(cell, buf)
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, WA_MINUTES)
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

# Naming, headings and grading wording per instrument. The KEY rule: whichever instrument the
# course's original assessment uses, keep using it — do not migrate CS ↔ PP.
INSTRUMENTS = {
    "PP": dict(cover="Practical Performance (PP)",
               heading="Practical Performance Assessment",
               section="Practical Problem",
               q_file="PP Assessment", a_file="Answer to PP Assessment",
               grade="Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them."),
    "CS": dict(cover="Case Study (CS)",
               heading="Case Study Assessment",
               section="Case Study",
               q_file="Case Study (CS)", a_file="Answer to Case Study (CS)",
               grade="Candidate has successfully completed all the tasks for the Case Study and is able to "
                     "explain the overall functions and features used to achieve these tasks."),
}

def build_practical(answers):
    inst = INSTRUMENTS[INSTRUMENT]
    doc = base_doc()
    kind = inst["cover"] + (" — Answer Key" if answers else "")
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, ("Answers to " + inst["heading"]) if answers else inst["heading"],
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, PRACTICAL_MINUTES)
        grading(doc, inst["grade"])
        page_break(doc)
    para(doc, inst["section"], size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    for label, crit, prompt, cap, pts in PRACTICAL:
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
    suffix = A_VER if answers else Q_VER
    name = f"{inst['a_file'] if answers else inst['q_file']} - {TITLE} - {suffix}.docx"
    finish(doc, os.path.join(OUT, name))

def _codes(criterion):
    """'A2, A3' -> ['A2','A3']"""
    return [c.strip() for c in str(criterion).split(",") if c.strip()]


def check_coverage():
    """Fail the build if any declared Ability/Knowledge item is not assessed anywhere,
    or if a question cites a code that is not in the competency unit."""
    wa_map, pp_map = {}, {}
    for i, item in enumerate(WRITTEN, 1):
        for c in _codes(item[0]):
            wa_map.setdefault(c, []).append(f"WA Q{i}")
    for item in PRACTICAL:
        label, crit = item[0], item[1]
        for c in _codes(crit):
            pp_map.setdefault(c, []).append(f"{INSTRUMENT} {label}")

    problems = []
    for code in KNOWLEDGE:
        if code not in wa_map:
            problems.append(f"  {code} is declared in the competency unit but NO Written Assessment question tests it.")
    for code in ABILITIES:
        if code not in pp_map:
            problems.append(f"  {code} is declared in the competency unit but NO {INSTRUMENTS[INSTRUMENT]['heading']} task tests it.")
    for code in wa_map:
        if code not in KNOWLEDGE:
            problems.append(f"  WA cites {code}, which is not in KNOWLEDGE {KNOWLEDGE}.")
    for code in pp_map:
        if code not in ABILITIES:
            problems.append(f"  {INSTRUMENT} cites {code}, which is not in ABILITIES {ABILITIES}.")

    print("\nCoverage map")
    for code in KNOWLEDGE:
        print(f"  {code:<4} → {', '.join(wa_map.get(code, [])) or '*** NOT ASSESSED ***'}")
    for code in ABILITIES:
        print(f"  {code:<4} → {', '.join(pp_map.get(code, [])) or '*** NOT ASSESSED ***'}")

    if problems:
        raise SystemExit("\nCOVERAGE CHECK FAILED — every A and K must be assessed:\n"
                         + "\n".join(problems))
    print("  Coverage OK — every ability and knowledge item is assessed.\n")


if __name__ == "__main__":
    print("Building WSQ assessment set…")
    check_coverage()
    build_wa(answers=False); build_wa(answers=True)
    build_practical(answers=False); build_practical(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · {INSTRUMENT}: {len(PRACTICAL)} tasks.")
